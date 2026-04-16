from pathlib import Path
import re

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

BASE = Path('.')
MD_PATH = BASE / 'project_report.md'
DOCX_PATH = BASE / 'project_report.docx'


def set_run_font(run, name='Times New Roman', size=12, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.strip())
    set_run_font(run, size=14, bold=True)


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=12)


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_run_font(run, size=12)


def add_code(doc, lines):
    for ln in lines:
        p = doc.add_paragraph()
        run = p.add_run(ln.rstrip('\n'))
        set_run_font(run, name='Consolas', size=10)


def parse_image(line):
    m = re.match(r'!\[.*?\]\((.*?)\)', line.strip())
    if not m:
        return None
    return m.group(1).strip()


def is_table_sep(line):
    s = line.strip().replace('|', '').replace('-', '').replace(':', '').replace(' ', '')
    return s == ''


def build_table(doc, table_lines):
    # Remove empty boundaries and split by '|'
    rows = []
    for ln in table_lines:
        ln = ln.strip()
        if not ln.startswith('|'):
            return False
        cols = [c.strip() for c in ln.strip('|').split('|')]
        rows.append(cols)

    if len(rows) < 2:
        return False

    # second row is separator in markdown
    if not is_table_sep(table_lines[1]):
        return False

    header = rows[0]
    data_rows = rows[2:]
    if not data_rows:
        data_rows = []

    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        r = hdr_cells[i].paragraphs[0].add_run(h)
        set_run_font(r, size=12, bold=True)

    for dr in data_rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(dr):
            r = row_cells[i].paragraphs[0].add_run(val)
            set_run_font(r, size=12)

    doc.add_paragraph('')
    return True


def main():
    if not MD_PATH.exists():
        raise FileNotFoundError('project_report.md not found')

    lines = MD_PATH.read_text(encoding='utf-8').splitlines()

    doc = Document()

    # Set default normal style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal_style.font.size = Pt(12)

    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                add_code(doc, code_lines)
                doc.add_paragraph('')
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            doc.add_paragraph('')
            i += 1
            continue

        # Table detection block
        if line.strip().startswith('|'):
            j = i
            tlines = []
            while j < len(lines) and lines[j].strip().startswith('|'):
                tlines.append(lines[j])
                j += 1
            if build_table(doc, tlines):
                i = j
                continue

        if line.startswith('# '):
            add_heading(doc, line[2:])
        elif line.startswith('## '):
            add_heading(doc, line[3:])
        elif line.startswith('### '):
            add_heading(doc, line[4:])
        elif line.startswith('#### '):
            add_heading(doc, line[5:])
        elif line.startswith('- '):
            add_bullet(doc, line[2:])
        elif re.match(r'^\d+\.\s+', line):
            p = doc.add_paragraph(style='List Number')
            txt = re.sub(r'^\d+\.\s+', '', line)
            r = p.add_run(txt)
            set_run_font(r, size=12)
        else:
            img_path = parse_image(line)
            if img_path:
                p = BASE / img_path
                if p.exists():
                    doc.add_picture(str(p), width=Inches(6.2))
                    doc.add_paragraph('')
                else:
                    add_paragraph(doc, f'[Missing image: {img_path}]')
            else:
                add_paragraph(doc, line)

        i += 1

    doc.save(str(DOCX_PATH))
    print(f'Generated {DOCX_PATH}')


if __name__ == '__main__':
    main()
